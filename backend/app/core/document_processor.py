import os
import pandas as pd
from pathlib import Path
from typing import List
from llama_index.core import Document
from llama_index.core.node_parser import SentenceSplitter

class DocumentProcessor:
    """Process different document types into LlamaIndex Documents"""
    
    # Text splitter for chunking large documents
    text_splitter = SentenceSplitter(
        chunk_size=1024,
        chunk_overlap=200
    )
    
    @staticmethod
    def process_pdf(file_path: str) -> List[Document]:
        """Process PDF files page by page with pdfplumber"""
        try:
            import pdfplumber
            
            documents = []
            filename = Path(file_path).name
            all_text = []
            
            with pdfplumber.open(file_path) as pdf:
                print(f"  PDF has {len(pdf.pages)} pages")
                
                for page_num, page in enumerate(pdf.pages, 1):
                    text = page.extract_text()
                    
                    if text and len(text.strip()) > 50:
                        clean_text = text.replace('\x00', '').strip()
                        all_text.append(f"\n--- Page {page_num} ---\n{clean_text}")
            
            # Combine all pages
            combined_text = "\n".join(all_text)
            
            # Split into chunks
            text_chunks = DocumentProcessor.text_splitter.split_text(combined_text)
            
            print(f"  Split into {len(text_chunks)} chunks")
            
            # Create documents for each chunk
            for i, chunk in enumerate(text_chunks):
                doc = Document(
                    text=chunk,
                    metadata={
                        "source_type": "pdf",
                        "filename": filename,
                        "chunk_id": i,
                        "total_chunks": len(text_chunks)
                    }
                )
                documents.append(doc)
            
            print(f"PDF: Extracted {len(documents)} chunks")
            return documents
            
        except ImportError:
            print("  pdfplumber not installed, using default reader with chunking")
            return DocumentProcessor._fallback_pdf_processing(file_path)
            
        except Exception as e:
            print(f"Error with pdfplumber: {e}")
            return DocumentProcessor._fallback_pdf_processing(file_path)
    
    @staticmethod
    def _fallback_pdf_processing(file_path: str) -> List[Document]:
        """Fallback PDF processing using SimpleDirectoryReader"""
        from llama_index.core import SimpleDirectoryReader
        
        try:
            docs = SimpleDirectoryReader(
                input_files=[file_path],
                filename_as_id=True
            ).load_data()
            
            all_docs = []
            filename = Path(file_path).name
            
            for doc in docs:
                clean_text = doc.text.replace('\x00', '').strip() if doc.text else ""
                
                if clean_text and len(clean_text) > 100:
                    # Split large document into chunks
                    text_chunks = DocumentProcessor.text_splitter.split_text(clean_text)
                    
                    for i, chunk in enumerate(text_chunks):
                        new_doc = Document(
                            text=chunk,
                            metadata={
                                "source_type": "pdf",
                                "filename": filename,
                                "chunk_id": i,
                                "total_chunks": len(text_chunks)
                            }
                        )
                        all_docs.append(new_doc)
            
            print(f"PDF: Extracted {len(all_docs)} chunks (fallback)")
            return all_docs
            
        except Exception as e:
            print(f"Error in fallback PDF processing: {e}")
            return []
    
    @staticmethod
    def process_word(file_path: str) -> List[Document]:
        """Process Word documents with python-docx"""
        try:
            from docx import Document as DocxDocument
            
            filename = Path(file_path).name
            
            # Read Word document
            docx = DocxDocument(file_path)
            
            # Extract all paragraphs
            paragraphs = []
            for para in docx.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append(text)
            
            # Extract tables
            for table in docx.tables:
                table_text = []
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells])
                    if row_text.strip():
                        table_text.append(row_text)
                
                if table_text:
                    paragraphs.append("\n".join(table_text))
            
            # Combine all text
            combined_text = "\n\n".join(paragraphs)
            
            if not combined_text or len(combined_text) < 50:
                print("  Word document appears empty")
                return []
            
            # Split into chunks
            text_chunks = DocumentProcessor.text_splitter.split_text(combined_text)
            
            print(f"  Word doc has {len(paragraphs)} paragraphs, split into {len(text_chunks)} chunks")
            
            # Create documents
            documents = []
            for i, chunk in enumerate(text_chunks):
                doc = Document(
                    text=chunk,
                    metadata={
                        "source_type": "word",
                        "filename": filename,
                        "chunk_id": i,
                        "total_chunks": len(text_chunks)
                    }
                )
                documents.append(doc)
            
            print(f"Word: Extracted {len(documents)} chunks")
            return documents
            
        except ImportError:
            print("  python-docx not installed, using default reader with chunking")
            return DocumentProcessor._fallback_word_processing(file_path)
            
        except Exception as e:
            print(f"Error with python-docx: {e}")
            return DocumentProcessor._fallback_word_processing(file_path)
    
    @staticmethod
    def _fallback_word_processing(file_path: str) -> List[Document]:
        """Fallback Word processing using SimpleDirectoryReader"""
        from llama_index.core import SimpleDirectoryReader
        
        try:
            docs = SimpleDirectoryReader(
                input_files=[file_path],
                filename_as_id=True
            ).load_data()
            
            all_docs = []
            filename = Path(file_path).name
            
            for doc in docs:
                clean_text = doc.text.replace('\x00', '').strip() if doc.text else ""
                
                if clean_text and len(clean_text) > 100:
                    # Split into chunks
                    text_chunks = DocumentProcessor.text_splitter.split_text(clean_text)
                    
                    for i, chunk in enumerate(text_chunks):
                        new_doc = Document(
                            text=chunk,
                            metadata={
                                "source_type": "word",
                                "filename": filename,
                                "chunk_id": i,
                                "total_chunks": len(text_chunks)
                            }
                        )
                        all_docs.append(new_doc)
            
            print(f"Word: Extracted {len(all_docs)} chunks (fallback)")
            return all_docs
            
        except Exception as e:
            print(f"Error in fallback Word processing: {e}")
            return []
    
    @staticmethod
    def process_excel(file_path: str) -> List[Document]:
        """Process Excel files with smart table parsing"""
        documents = []
        filename = Path(file_path).name
        
        try:
            df = pd.read_excel(file_path, sheet_name='Sheet1', header=None)
            
            current_product = None
            doc_count = 0
            
            for idx in range(len(df)):
                cell_val = df.iloc[idx, 0]
                
                # Detect product headers
                if pd.notna(cell_val) and 'European' in str(cell_val) and 'per ton' in str(cell_val):
                    current_product = str(cell_val).replace(' (€ per ton)', '').strip()
                    print(f"  Found product: {current_product}")
                
                # Detect year rows
                if pd.notna(df.iloc[idx, 2]) and df.iloc[idx, 2] in [2023.0, 2024.0, 2025.0]:
                    year = int(df.iloc[idx, 2])
                    
                    months = ['February', 'March', 'April', 'May', 'June', 'July', 
                             'August', 'September', 'October', 'November', 'December']
                    
                    for row_offset in range(1, 7):
                        if idx + row_offset >= len(df):
                            break
                        
                        row_type = df.iloc[idx + row_offset, 0]
                        if pd.isna(row_type) or row_type == '':
                            break
                        
                        trade_type = str(row_type).split()[-1] if pd.notna(row_type) else "Unknown"
                        
                        for month_idx, month in enumerate(months):
                            col_idx = 3 + month_idx
                            value = df.iloc[idx + row_offset, col_idx]
                            
                            if pd.notna(value) and value != 0:
                                text = f"""
Product: {current_product}
Country/Type: {row_type}
Trade Type: {trade_type}
Month: {month}
Year: {year}
Price: €{value:.2f} per ton

Question patterns this answers:
- What was the {current_product} {trade_type.lower()} price in {month} {year}?
- {row_type} {month} {year} {current_product} price
- {month} {year} {current_product} {trade_type.lower()}

Data: {current_product} {row_type} {trade_type} {month} {year} = €{value:.2f} per ton
"""
                                
                                doc = Document(
                                    text=text,
                                    metadata={
                                        "source_type": "excel",
                                        "filename": filename,
                                        "product": current_product,
                                        "type": str(row_type),
                                        "trade_type": trade_type,
                                        "month": month,
                                        "year": year,
                                        "value": float(value)
                                    }
                                )
                                documents.append(doc)
                                doc_count += 1
            
            print(f"Excel: Created {doc_count} searchable entries")
        
        except Exception as e:
            print(f"Error processing Excel: {e}")
            import traceback
            traceback.print_exc()
        
        return documents
    
    @staticmethod
    def process_document(file_path: str) -> List[Document]:
        """Route to appropriate processor based on file type"""
        ext = Path(file_path).suffix.lower()
        
        print(f"Processing {ext} file: {Path(file_path).name}")
        
        if ext == '.pdf':
            return DocumentProcessor.process_pdf(file_path)
        elif ext in ['.docx', '.doc']:
            return DocumentProcessor.process_word(file_path)
        elif ext in ['.xlsx', '.xls']:
            return DocumentProcessor.process_excel(file_path)
        else:
            raise ValueError(f"Unsupported file type: {ext}")